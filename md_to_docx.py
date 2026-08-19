#!/usr/bin/env python3
from __future__ import annotations

import argparse
import base64
import io
import json
import re
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PRIMARY = "0C1A32"
ACCENT = "FF5722"
LIGHT_BG = "F6F8FB"
TABLE_HEADER_BG = "0C1A32"
TABLE_BORDER = "D8DEE9"
BODY_FONT = "PingFang SC"
WEST_FONT = "Arial"


def set_run_font(run, size: int | float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = WEST_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size: int | float = 10.5, color: str = "1F2937"):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), TABLE_BORDER)


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()


def write_cell(cell, text: str, header: bool = False):
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, size=9.5, bold=header, color="FFFFFF" if header else "1F2937")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    if header:
        set_cell_shading(cell, TABLE_HEADER_BG)
    else:
        set_cell_shading(cell, "FFFFFF")


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\{width=[^}]+\}", "", text)
    text = re.sub(r"<[^>]+>", "", text)
    return text.strip()


def parse_inline_tokens(text: str) -> list[tuple[str, bool, bool]]:
    # 替换简单 HTML 标签
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</?(p|div|span)[^>]*>", "", text, flags=re.IGNORECASE)
    text = text.replace("<strong>", "**").replace("</strong>", "**")
    text = re.sub(r"\{width=[^}]+\}", "", text)
    
    # 匹配 **bold** 与 `code`
    pattern = r"(\*\*.+?\*\*|`[^`]+`)"
    tokens = re.split(pattern, text)
    result = []
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) >= 4:
            result.append((token[2:-2], True, False))
        elif token.startswith("`") and token.endswith("`") and len(token) >= 2:
            result.append((token[1:-1], False, True))
        else:
            result.append((token, False, False))
    return result


def add_formatted_text(paragraph, text: str, default_size=10.5, default_color="1F2937", force_bold=False):
    tokens = parse_inline_tokens(text)
    for content, is_bold, is_code in tokens:
        # 处理可能的换行
        lines = content.split("\n")
        for i, line_text in enumerate(lines):
            if i > 0:
                paragraph.add_run().add_break()
            if not line_text:
                continue
            run = paragraph.add_run(line_text)
            if is_code:
                run.font.name = "Courier New"
                set_run_font(run, size=default_size - 0.5, bold=force_bold, color="334155")
            else:
                set_run_font(run, size=default_size, bold=is_bold or force_bold, color=default_color)


def add_paragraph_with_inline(paragraph, text: str, size=10.5, color="1F2937", bold=False):
    add_formatted_text(paragraph, text, default_size=size, default_color=color, force_bold=bold)


def set_paragraph_callout_style(paragraph, fill_color="F8FAFC", border_color="0C1A32"):
    pPr = paragraph._p.get_or_add_pPr()
    
    # 背景着色
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)
    
    # 左侧加粗边框
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24') # 3pt 边框宽度
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), border_color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_quote_block(doc: Document, quote_lines: list[str]):
    if not quote_lines:
        return
    
    for idx, line_text in enumerate(quote_lines):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.6)
        paragraph.paragraph_format.right_indent = Cm(0.4)
        
        # 首尾段落控制上下边距
        before = 8 if idx == 0 else 2
        after = 8 if idx == len(quote_lines) - 1 else 2
        set_paragraph_spacing(paragraph, before=before, after=after, line=1.35)
        set_paragraph_callout_style(paragraph, fill_color="F8FAFC", border_color="0C1A32")
        
        # 检查是否为引用块内部的列表项 (如 > 1. xxx 或 > - xxx)
        numbered_match = re.match(r"^(\d+)\.\s+(.+)$", line_text)
        bullet_match = re.match(r"^[-*]\s+(.+)$", line_text)
        
        if numbered_match:
            num_str = numbered_match.group(1)
            content = numbered_match.group(2)
            prefix_run = paragraph.add_run(f"{num_str}. ")
            set_run_font(prefix_run, size=10.5, bold=True, color=PRIMARY)
            add_formatted_text(paragraph, content, default_size=10.5, default_color="1F2937")
        elif bullet_match:
            content = bullet_match.group(1)
            prefix_run = paragraph.add_run("• ")
            set_run_font(prefix_run, size=10.5, bold=True, color=ACCENT)
            add_formatted_text(paragraph, content, default_size=10.5, default_color="1F2937")
        else:
            add_formatted_text(paragraph, line_text, default_size=10.5, default_color="1F2937")


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = WEST_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1F2937")

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = WEST_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.color.rgb = RGBColor.from_string(PRIMARY)
        style.font.bold = True

    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 3"].font.size = Pt(12)


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.35):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bookmark(paragraph, name: str, bookmark_id: int):
    tag_start = OxmlElement('w:bookmarkStart')
    tag_start.set(qn('w:id'), str(bookmark_id))
    tag_start.set(qn('w:name'), name)
    tag_end = OxmlElement('w:bookmarkEnd')
    tag_end.set(qn('w:id'), str(bookmark_id))
    paragraph._p.append(tag_start)
    paragraph._p.append(tag_end)


def add_hyperlink_to_bookmark(paragraph, text: str, bookmark_name: str, level: int = 1):
    indent_cm = 0.2 if level == 1 else (0.6 if level == 2 else 1.0)
    paragraph.paragraph_format.left_indent = Cm(indent_cm)
    set_paragraph_spacing(paragraph, after=4, line=1.25)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    hyperlink.set(qn('w:history'), '1')
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # 字体与颜色
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), WEST_FONT)
    rFonts.set(qn('w:eastAsia'), BODY_FONT)
    rPr.append(rFonts)
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0284C7')
    rPr.append(color)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21' if level == 1 else '20')
    rPr.append(sz)
    
    if level == 1:
        b = OxmlElement('w:b')
        rPr.append(b)
        
    new_run.append(rPr)
    
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_native_toc_field(paragraph):
    set_paragraph_spacing(paragraph, before=6, after=12)
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)


def write_cell(cell, text: str, header: bool = False, zebra: bool = False):
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 允许单元格中解析行内粗体与 <br> 换行
    add_formatted_text(paragraph, text, default_size=9.5, default_color="FFFFFF" if header else "1F2937", force_bold=header)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
    
    if header:
        set_cell_shading(cell, TABLE_HEADER_BG)
    elif zebra:
        set_cell_shading(cell, "F8FAFC")
    else:
        set_cell_shading(cell, "FFFFFF")


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.autofit = True
    set_table_borders(table)
    
    for row_index, row in enumerate(rows):
        is_header = (row_index == 0)
        is_zebra = (row_index % 2 == 1 and not is_header)
        for col_index in range(width):
            text = row[col_index] if col_index < len(row) else ""
            write_cell(table.rows[row_index].cells[col_index], text, header=is_header, zebra=is_zebra)
            
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=10)


def add_heading(doc: Document, text: str, level: int, bookmark_name: str | None = None, bookmark_id: int = 0):
    clean_title = clean_inline(text)
    
    if level == 1:
        # 大标题 (H1)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(paragraph, before=6, after=14, line=1.15)
        run = paragraph.add_run(clean_title)
        set_run_font(run, size=22, bold=True, color=PRIMARY)

        if bookmark_name:
            add_bookmark(paragraph, bookmark_name, bookmark_id)
        return

    # 二级与三级标题
    paragraph = doc.add_heading(clean_title, level=min(level, 3))
    
    if level == 2:
        # 二级标题 (H2)：加上边框装饰与背景衬底
        set_paragraph_spacing(paragraph, before=18, after=8, line=1.25)
        set_paragraph_callout_style(paragraph, fill_color="F1F5F9", border_color=PRIMARY)
        paragraph.paragraph_format.left_indent = Cm(0.2)
        for run in paragraph.runs:
            set_run_font(run, size=14.5, bold=True, color=PRIMARY)
    else:
        # 三级标题 (H3)
        set_paragraph_spacing(paragraph, before=12, after=6, line=1.2)
        for run in paragraph.runs:
            set_run_font(run, size=12, bold=True, color=ACCENT)

    if bookmark_name:
        add_bookmark(paragraph, bookmark_name, bookmark_id)


def add_body_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=9, line=1.55)
    add_paragraph_with_inline(paragraph, text, size=11)


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.55)
    paragraph.paragraph_format.first_line_indent = Cm(-0.28)
    set_paragraph_spacing(paragraph, after=6, line=1.5)
    bullet = paragraph.add_run("• ")
    set_run_font(bullet, size=11, bold=False, color=ACCENT)
    add_formatted_text(paragraph, text, default_size=11, default_color="1F2937")


def add_numbered(doc: Document, number: str, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.55)
    paragraph.paragraph_format.first_line_indent = Cm(-0.4)
    set_paragraph_spacing(paragraph, after=7, line=1.52)
    prefix = paragraph.add_run(f"{number}. ")
    set_run_font(prefix, size=11, bold=True, color=PRIMARY)
    add_formatted_text(paragraph, text, default_size=11, default_color="1F2937")


def add_quote(doc: Document, text: str):
    add_quote_block(doc, [text])


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    index = start
    while index < len(lines):
        line = lines[index].strip()
        if not line.startswith("|") or not line.endswith("|"):
            break
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            index += 1
            continue
        rows.append(cells)
        index += 1
    return rows, index


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.autofit = True
    set_table_borders(table)
    for row_index, row in enumerate(rows):
        for col_index in range(width):
            text = row[col_index] if col_index < len(row) else ""
            write_cell(table.rows[row_index].cells[col_index], text, header=row_index == 0)
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=8)


def add_image(doc: Document, md_path: Path, line: str):
    match = re.match(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)(?P<attrs>\{[^}]+\})?", line.strip())
    if not match:
        add_body_paragraph(doc, line)
        return
    alt = match.group("alt").strip()
    image_path = (md_path.parent / match.group("path")).resolve()
    if not image_path.exists():
        add_body_paragraph(doc, f"[图片缺失] {alt}：{image_path}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=8, after=4)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(15.8))
    if alt:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(caption, after=10)
        caption_run = caption.add_run(alt)
        set_run_font(caption_run, size=9, color="64748B")


def sanitize_mermaid_code(code: str) -> str:
    lines = code.splitlines()
    sanitized = []
    is_sequence = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("sequenceDiagram"):
            is_sequence = True
            sanitized.append(line)
            continue
        if is_sequence:
            m_actor = re.match(r"^(\s*actor\s+)([^\"'\s]+)(.*)$", line)
            if m_actor:
                line = f'{m_actor.group(1)}"{m_actor.group(2)}"{m_actor.group(3)}'
            m_part = re.match(r"^(\s*participant\s+[^\"'\s]+\s+as\s+)([^\"'].*)$", line)
            if m_part and not m_part.group(2).startswith('"'):
                line = f'{m_part.group(1)}"{m_part.group(2).strip()}"'
        sanitized.append(line)
    return "\n".join(sanitized)


import json

def add_mermaid_image(doc: Document, mermaid_code: str):
    png_data = None
    svg_data = None

    # 1. 抓取 SVG 矢量数据与 PNG 高清数据
    try:
        graph_dict = {"code": mermaid_code.strip(), "mermaid": {"theme": "default"}}
        base64_str = base64.b64encode(json.dumps(graph_dict).encode("utf-8")).decode("utf-8")
        
        svg_url = f"https://mermaid.ink/svg/{base64_str}"
        png_url = f"https://mermaid.ink/img/{base64_str}"

        req_svg = urllib.request.Request(svg_url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req_svg, timeout=15) as resp:
            svg_data = resp.read()

        req_png = urllib.request.Request(png_url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req_png, timeout=15) as resp:
            png_data = resp.read()
    except Exception as e:
        print(f"mermaid.ink fetch failed: {e}")

    # 备用 kroki
    if not svg_data or not png_data:
        try:
            req_kroki = urllib.request.Request(
                "https://kroki.io/mermaid/png",
                data=mermaid_code.strip().encode("utf-8"),
                headers={"Content-Type": "text/plain; charset=utf-8", "User-Agent": "Mozilla/5.0"}
            )
            with urllib.request.urlopen(req_kroki, timeout=15) as resp:
                png_data = resp.read()
        except Exception as e:
            print(f"kroki fetch failed: {e}")

    # 绘制高清晰度流程图节点
    if png_data:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, before=10, after=10)
        run = paragraph.add_run()
        image_stream = io.BytesIO(png_data)
        run.add_picture(image_stream, width=Cm(15.8))

        # 附加矢量 SVG 结构
        if svg_data:
            try:
                from docx.opc.constants import RELATIONSHIP_TYPE
                from docx.opc.packuri import PackURI
                from docx.opc.part import Part
                from docx.oxml import parse_xml

                part_count = len(doc.part.package.parts)
                svg_pack_uri = PackURI(f"/word/media/image_svg_{part_count}.svg")
                svg_part = Part(svg_pack_uri, "image/svg+xml", svg_data, doc.part.package)
                svg_rId = doc.part.relate_to(svg_part, RELATIONSHIP_TYPE.IMAGE)

                blip = run._r.xpath(".//a:blip")[0]
                svg_blip_xml = f'<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:embed="{svg_rId}"/>'
                blip.append(parse_xml(svg_blip_xml))
            except Exception as svg_err:
                print(f"Failed to attach SVG blip: {svg_err}")
    else:
        add_quote(doc, f"[Mermaid 流程图代码]:\n{mermaid_code}")


def add_code_block(doc: Document, code_text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    set_paragraph_spacing(paragraph, before=6, after=8, line=1.15)
    run = paragraph.add_run(code_text)
    set_run_font(run, size=9.5, color="334155")
    run.font.name = "Courier New"


def add_footer(doc: Document):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("爻鉴 · 数字内容资产整理工具链说明")
    set_run_font(run, size=8, color="64748B")


def convert(md_path: Path, output_path: Path):
    doc = Document()
    configure_styles(doc)
    add_footer(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    # 预扫描所有标题，建立索引字典
    collected_headings = []
    heading_counter = 0
    heading_map = {}
    for idx, raw in enumerate(lines):
        l = raw.strip()
        m1 = re.match(r"^(#{1,6})\s+(.+)$", l)
        m2 = re.match(r"^<h([1-6])>(.+?)</h\1>$", l)
        if m1:
            level = len(m1.group(1))
            text = clean_inline(m1.group(2))
            heading_counter += 1
            bname = f"_TocHeading_{heading_counter}"
            collected_headings.append((text, level, bname, heading_counter, idx))
            heading_map[idx] = (bname, heading_counter)
        elif m2:
            level = int(m2.group(1))
            text = clean_inline(m2.group(2))
            heading_counter += 1
            bname = f"_TocHeading_{heading_counter}"
            collected_headings.append((text, level, bname, heading_counter, idx))
            heading_map[idx] = (bname, heading_counter)

    index = 0
    toc_inserted = False

    # 闭包：插入超链接目录
    def insert_toc():
        nonlocal toc_inserted
        if toc_inserted or not collected_headings:
            return
        toc_inserted = True

        toc_heading = doc.add_paragraph()
        set_paragraph_spacing(toc_heading, before=10, after=10)
        toc_run = toc_heading.add_run("目录 / Table of Contents")
        set_run_font(toc_run, size=14, bold=True, color=PRIMARY)

        for h_text, h_level, h_bname, _, _ in collected_headings:
            if h_level <= 3:
                p = doc.add_paragraph()
                add_hyperlink_to_bookmark(p, h_text, h_bname, level=h_level)

        p_native = doc.add_paragraph()
        add_native_toc_field(p_native)

        sep_p = doc.add_paragraph()
        set_paragraph_spacing(sep_p, after=14)

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue

        if re.fullmatch(r"-{3,}", line):
            index += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            if not toc_inserted:
                insert_toc()
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            binfo = heading_map.get(index)
            bname, bid = binfo if binfo else (None, 0)
            
            add_heading(doc, text, level, bookmark_name=bname, bookmark_id=bid)
            
            # 如果是一级大标题 (H1)，目录插在 H1 之后；如果是 H2/H3，且目录尚未插入，直接插在最前方
            if level == 1 and not toc_inserted:
                insert_toc()
            elif level > 1 and not toc_inserted:
                insert_toc()
                
            index += 1
            continue

        if line.startswith("!["):
            if not toc_inserted:
                insert_toc()
            add_image(doc, md_path, line)
            index += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            if not toc_inserted:
                insert_toc()
            add_bullet(doc, bullet.group(1))
            index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", line)
        if numbered:
            if not toc_inserted:
                insert_toc()
            add_numbered(doc, numbered.group(1), numbered.group(2))
            index += 1
            continue

        if line.startswith("> "):
            if not toc_inserted:
                insert_toc()
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith("> "):
                quote_lines.append(lines[index].strip()[2:].strip())
                index += 1
            add_quote_block(doc, quote_lines)
            continue

        if line.startswith("```"):
            if not toc_inserted:
                insert_toc()
            if line.startswith("```mermaid"):
                mermaid_lines = []
                index += 1
                while index < len(lines) and not lines[index].strip().startswith("```"):
                    mermaid_lines.append(lines[index])
                    index += 1
                if index < len(lines) and lines[index].strip().startswith("```"):
                    index += 1
                add_mermaid_image(doc, "\n".join(mermaid_lines))
                continue
            else:
                code_lines = []
                index += 1
                while index < len(lines) and not lines[index].strip().startswith("```"):
                    code_lines.append(lines[index])
                    index += 1
                if index < len(lines) and lines[index].strip().startswith("```"):
                    index += 1
                add_code_block(doc, "\n".join(code_lines))
                continue

        html_heading = re.match(r"^<h([1-6])>(.+?)</h\1>$", line)
        if html_heading:
            level = int(html_heading.group(1))
            text = html_heading.group(2)
            binfo = heading_map.get(index)
            bname, bid = binfo if binfo else (None, 0)
            
            add_heading(doc, text, level, bookmark_name=bname, bookmark_id=bid)
            if not toc_inserted:
                insert_toc()
            index += 1
            continue

        # 剥离简单的 <p> 标签包围
        if line.startswith("<p>") and line.endswith("</p>"):
            line = line[3:-4]
            # 如果里面还有标签，会在 clean_inline 里被剥离
            
        # 如果是 HTML 注释，直接跳过
        if line.startswith("<!--") and line.endswith("-->"):
            index += 1
            continue

        if not toc_inserted:
            insert_toc()

        add_body_paragraph(doc, line)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Convert project Markdown to a structured Word document.")
    parser.add_argument("input", type=Path)
    parser.add_argument("-o", "--output", type=Path, required=True)
    args = parser.parse_args()
    convert(args.input.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
